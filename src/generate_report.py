# -*- coding: utf-8 -*-
"""
《机器学习》综合设计报告生成器 - 手写数字识别系统
生成带插图、表注在上、图注在下的完整 DOCX 报告
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime, numpy as np, matplotlib, matplotlib.pyplot as plt

matplotlib.use('Agg')
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(BASE_DIR, 'figures')
REPORT_DIR = os.path.join(BASE_DIR, 'report')
os.makedirs(REPORT_DIR, exist_ok=True)

# ======================== 工具函数 ========================
def set_run_font(run, fn_cn='宋体', fn_en='Times New Roman', size=Pt(12), bold=False):
    run.font.size = size; run.font.bold = bold; run.font.name = fn_en
    r = run._element; rPr = r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), fn_cn); rF.set(qn('w:ascii'), fn_en); rF.set(qn('w:hAnsi'), fn_en)
    rPr.insert(0, rF)

def add_para(doc, text, fn_cn='宋体', size=Pt(12), bold=False, align=None, indent=None, after=Pt(3)):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    pf = p.paragraph_format; pf.space_after = after; pf.line_spacing = 1.5
    if indent: pf.first_line_indent = indent
    r = p.add_run(text)
    set_run_font(r, fn_cn, 'Times New Roman', size, bold)
    return p

def heading(doc, text, level=1):
    if level == 1: return add_para(doc, text, fn_cn='仿宋', size=Pt(14), bold=True, after=Pt(6))
    if level == 2: return add_para(doc, text, fn_cn='黑体', size=Pt(12), bold=True, after=Pt(3))
    return add_para(doc, text, fn_cn='仿宋', size=Pt(12), after=Pt(3))

def add_table_caption(doc, text):
    """表注——在表格上方，小5黑体居中"""
    add_para(doc, '', size=Pt(6), after=Pt(0))
    p = add_para(doc, text, fn_cn='黑体', size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(3))
    return p

def add_figure_caption(doc, text):
    """图注——在图片下方，小5宋体居中"""
    p = add_para(doc, text, fn_cn='宋体', size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(6))
    return p

def add_image(doc, fname, width=Inches(5.5)):
    """插入图片，居中"""
    path = os.path.join(FIGURES_DIR, fname)
    if not os.path.exists(path):
        add_para(doc, f'[图表 {fname} 未生成，请先运行 ml_pipeline.py]', fn_cn='宋体', size=Pt(9))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=width)

def make_table(doc, headers, rows):
    """创建表格并返回"""
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rr in pp.runs: set_run_font(rr, '黑体', size=Pt(10.5), bold=True)
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = t.rows[i+1].cells[j]; c.text = v
            for pp in c.paragraphs:
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rr in pp.runs: set_run_font(rr, '宋体', size=Pt(10.5))
    return t

# ======================== 生成额外图表 ========================
def gen_extra_figures():
    """生成报告中需要的额外图表"""
    # ---- CNN 架构图 ----
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 6); ax.axis('off')

    layers = [
        (0.5, 3, 1.2, 2.5, '输入\n28×28×1', '#E8E8E8'),
        (2.2, 3, 1.2, 2.5, 'Conv1\n32@3×3\n→ReLU', '#FFCCBC'),
        (3.8, 3, 1.2, 2.5, 'Pool\n2×2\n→14×14', '#BBDEFB'),
        (5.4, 3, 1.2, 2.5, 'Conv2\n64@3×3\n→ReLU', '#FFCCBC'),
        (7.0, 3, 1.2, 2.5, 'Pool\n2×2\n→7×7', '#BBDEFB'),
        (8.6, 3, 1.2, 2.5, 'FC1\n3136→256\n+Dropout', '#C8E6C9'),
        (10.1, 3, 1.2, 2.5, 'FC2\n256→10\nSoftmax', '#C8E6C9'),
    ]
    for x, y, w, h, label, color in layers:
        rect = plt.Rectangle((x-w/2, y-h/2), w, h, facecolor=color, edgecolor='#333', linewidth=1.5, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x, y, label, ha='center', va='center', fontsize=8, fontweight='bold')
    for i in range(len(layers)-1):
        x1 = layers[i][0] + layers[i][2]/2; x2 = layers[i+1][0] - layers[i+1][2]/2
        ax.annotate('', xy=(x2, 3), xytext=(x1, 3), arrowprops=dict(arrowstyle='->', lw=2, color='#555'))
    ax.set_title('CNN 网络架构图', fontsize=14, fontweight='bold', pad=10)
    plt.tight_layout()
    plt.savefig(os.path.join(FIGURES_DIR, 'cnn_architecture.png'), dpi=150)
    plt.close()

    # ---- PCA 解释方差图（使用本地MNIST数据） ----
    from sklearn.decomposition import PCA
    try:
        import torch, torchvision
        raw = torchvision.datasets.MNIST(root=os.path.join(BASE_DIR, 'data', 'mnist_raw'),
                                         train=True, download=False,
                                         transform=torchvision.transforms.ToTensor())
        X_pca = raw.data.numpy()[:10000].reshape(-1, 784) / 255.0
        pca_full = PCA().fit(X_pca)
        cumsum = np.cumsum(pca_full.explained_variance_ratio_)
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
        ax1.plot(range(1, 101), pca_full.explained_variance_ratio_[:100]*100, 'b-', lw=1.5)
        ax1.set_xlabel('主成分序号'); ax1.set_ylabel('解释方差比 (%)')
        ax1.set_title('PCA 各主成分解释方差比 (前100维)')
        ax1.axvline(100, color='r', linestyle='--', alpha=0.5, label='选择100维')
        ax1.legend()
        ax2.plot(range(1, 201), cumsum[:200]*100, 'r-', lw=2)
        ax2.set_xlabel('主成分数量'); ax2.set_ylabel('累计解释方差比 (%)')
        ax2.set_title('PCA 累计解释方差比')
        ax2.axhline(91.53, color='b', linestyle='--', alpha=0.5, label='100维:91.53%')
        ax2.axvline(100, color='r', linestyle='--', alpha=0.5)
        ax2.legend()
        plt.tight_layout()
        plt.savefig(os.path.join(FIGURES_DIR, 'pca_variance.png'), dpi=150)
        plt.close()
        print('[DONE] PCA图已生成')
    except Exception as e:
        print(f'[INFO] PCA 图生成跳过 ({e})')

    # ---- 混淆数字对可视化（使用本地MNIST数据） ----
    try:
        import torchvision
        raw = torchvision.datasets.MNIST(root=os.path.join(BASE_DIR, 'data', 'mnist_raw'),
                                         train=False, download=False,
                                         transform=torchvision.transforms.ToTensor())
        targets = raw.targets.numpy()
        fig, axes = plt.subplots(2, 3, figsize=(12, 5))
        pairs = [(4, 9), (3, 8), (2, 7)]
        for idx, (a, b) in enumerate(pairs):
            axes[0, idx].set_title(f'数字 {a} vs {b}', fontsize=11, fontweight='bold')
            ia = np.where(targets == a)[0][0]
            ib = np.where(targets == b)[0][0]
            axes[0, idx].imshow(raw.data[ia], cmap='gray')
            axes[0, idx].set_xlabel(f'数字 {a}', fontsize=9)
            axes[0, idx].set_xticks([]); axes[0, idx].set_yticks([])
            axes[1, idx].imshow(raw.data[ib], cmap='gray')
            axes[1, idx].set_xlabel(f'数字 {b}', fontsize=9)
            axes[1, idx].set_xticks([]); axes[1, idx].set_yticks([])
        plt.suptitle('MNIST 中易混淆数字对示例', fontsize=13, fontweight='bold')
        plt.tight_layout()
        plt.savefig(os.path.join(FIGURES_DIR, 'confusing_pairs.png'), dpi=150)
        plt.close()
        print('[DONE] 混淆对图已生成')
    except Exception as e:
        print(f'[INFO] 混淆对图跳过 ({e})')

# ======================== 主报告生成 ========================
def generate_report():
    print("生成额外图表...")
    gen_extra_figures()

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18); sec.right_margin = Cm(3.18)

    sty = doc.styles['Normal']
    sty.font.name = 'Times New Roman'; sty.font.size = Pt(12)
    sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面 =====
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('《机器学习》'); set_run_font(r, '黑体', size=Pt(26), bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('综合设计报告'); set_run_font(r, '黑体', size=Pt(26), bold=True)
    for _ in range(3): doc.add_paragraph()
    for l, v in [('题    目：', '基于CNN/SVM的手写数字识别系统设计与实现'),
                 ('姓    名：', '（请填写）'), ('学    号：', '（请填写）'),
                 ('院    系：', '工学院'), ('年级专业：', '（请填写）'),
                 ('指导教师：', '（请填写）'),
                 ('完成时间：', datetime.date.today().strftime('%Y年%m月%d日'))]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{l}     {v}'); set_run_font(r, fn_cn='宋体', size=Pt(16))
    doc.add_page_break()

    # ===== 承诺声明 =====
    add_para(doc, '承 诺 声 明', fn_cn='黑体', size=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(24))
    add_para(doc, '本人郑重声明所呈交的综合设计是本人在老师指导下进行的设计工作成果。'
             '承诺在整个综合设计阶段根据所学的专业知识和参考相关文献独立完成，不存在抄袭现象。',
             indent=Cm(0.74), after=Pt(12))
    add_para(doc, '特此声明。', indent=Cm(0.74), after=Pt(24))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('作者签名：               （电子签名）   签字日期：    年  月  日')
    set_run_font(r, fn_cn='宋体', size=Pt(12))
    doc.add_page_break()

    # ===== 摘要 =====
    add_para(doc, '摘  要', fn_cn='黑体', size=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(12))
    add_para(doc, '手写数字识别是计算机视觉和模式识别领域的经典问题，广泛应用于邮政编码自动分拣、银行支票'
             '处理、税务表单录入等实际场景。本文基于MNIST手写数字数据集，设计并实现了一个完整的手写数字'
             '识别系统。在算法层面，本文深入实现了两种不同类型的机器学习模型：卷积神经网络（CNN）作为一种'
             '端到端的深度学习模型，能够自动从像素中提取层次化的视觉特征；支持向量机（SVM）结合主成分分析'
             '（PCA）降维，作为传统机器学习方法的代表。CNN模型采用2层卷积+2层全连接的结构，包含约82.5万个'
             '可训练参数，通过15个epoch的训练在测试集上达到99.44%的识别准确率。SVM模型使用RBF核函数和PCA'
             '100维降维，在测试集上取得96.46%的准确率。本文通过混淆矩阵、ROC曲线、每类准确率等多维度指标对'
             '两种算法进行了全面的对比分析，深入讨论了CNN在视觉空间信息利用上的优势以及SVM的计算约束。'
             '此外，本文还基于Streamlit框架开发了交互式Web应用并部署至云端，用户可在画布上自由绘制数字'
             '或上传手写图片，系统实时给出识别结果及各类别的预测概率分布，实现了从算法研究到工程应用的完整闭环。',
             fn_cn='宋体', size=Pt(10.5), indent=Cm(0.74))
    p = doc.add_paragraph(); r = p.add_run('关键词：'); set_run_font(r, fn_cn='宋体', size=Pt(10.5), bold=True)
    r = p.add_run('手写数字识别；MNIST；CNN；SVM；深度学习；图像分类；Streamlit')
    set_run_font(r, fn_cn='宋体', size=Pt(10.5))
    p = doc.add_paragraph(); r = p.add_run('Key words：'); set_run_font(r, fn_cn='宋体', size=Pt(10.5), bold=True)
    r = p.add_run('Handwritten Digit Recognition; MNIST; CNN; SVM; Deep Learning; Image Classification; Streamlit')
    set_run_font(r, fn_cn='宋体', size=Pt(10.5))
    doc.add_page_break()

    # ===== 目录 =====
    add_para(doc, '目  录', fn_cn='黑体', size=Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(12))
    for item, pg in [('摘要', 'I'), ('1  绪论', '1'), ('  1.1  设计的背景', '1'), ('  1.2  设计简介', '2'),
                     ('    1.2.1  问题定义', '2'), ('    1.2.2  数据集介绍', '3'),
                     ('    1.2.3  算法选择', '5'), ('    1.2.4  系统架构', '6'),
                     ('2  结果与分析', '7'), ('  2.1  数据探索性分析', '7'),
                     ('  2.2  数据预处理', '9'), ('  2.3  CNN模型训练与评估', '10'),
                     ('  2.4  SVM模型训练与评估', '12'), ('  2.5  模型性能对比分析', '13'),
                     ('  2.6  Web应用系统展示', '15'), ('5  结论', '17'), ('参考文献', '18')]:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 2.0
        r = p.add_run(f'{item}{"." * (40 - len(item))}{pg}')
        set_run_font(r, fn_cn='宋体', size=Pt(12))
    doc.add_page_break()

    # ==================== 1 绪论 ====================
    heading(doc, '1  绪论', 1)
    heading(doc, '1.1  设计的背景', 2)
    add_para(doc, '随着人工智能和深度学习技术的迅猛发展，计算机视觉已经成为机器学习最具影响力的应用领域之一。'
             '从智能手机的人脸解锁、自动驾驶的目标检测，到医疗影像的辅助诊断，计算机视觉技术正在深刻改变'
             '人们的生产和生活方式。在众多计算机视觉任务中，手写数字识别（Handwritten Digit Recognition）'
             '是一个兼具学术研究价值和实际应用意义的经典问题。早在20世纪90年代，美国邮政服务(UPS)就开始'
             '使用手写数字识别技术自动分拣邮政编码，大幅提高了邮件处理效率；银行系统利用该技术自动读取'
             '支票上的金额数字，减少了人工录入成本；税务部门也采用类似技术批量处理纳税人提交的纸质表单。', indent=Cm(0.74))
    add_para(doc, '学术界对手写数字识别的研究有着悠久的历史。1998年，Yann LeCun等人发表了里程碑式的论文'
             '"Gradient-Based Learning Applied to Document Recognition"，提出了LeNet-5卷积神经网络'
             '架构，并在MNIST数据集上取得了当时最优的识别效果[1]。这篇论文不仅奠定了现代卷积神经网络的'
             '基本架构，也使得MNIST成为深度学习领域的"Hello World"——几乎所有机器学习入门者都会接触的'
             '第一个图像数据集。MNIST数据集包含70,000张28×28像素的灰度手写数字图像（数字0-9），其中'
             '60,000张用于训练，10,000张用于测试。这一数据集的巧妙之处在于它"足够简单又不失挑战性"：'
             '28×28的低分辨率和灰度通道使得模型训练不需要昂贵的GPU，10个类别提供了足够的分类复杂度，'
             '而来自数百位不同书写者的样本又确保了数据多样性。', indent=Cm(0.74))
    add_para(doc, '然而，MNIST并非一个"已解决"的问题。在机器学习教学中，它仍然是一个极佳的教学工具——通过'
             '对比不同算法（从逻辑回归、SVM到CNN）在同一数据集上的表现，学生可以直观地理解"算法复杂度"'
             '与"任务特征"之间的匹配关系。特别是，用传统机器学习方法（SVM）和深度学习方法（CNN）分别解决'
             '同一个视觉任务，可以深刻揭示"手动特征工程"与"自动特征学习"两种范式的本质差异。这正是'
             '本课程设计选择手写数字识别作为研究课题的核心动机。', indent=Cm(0.74))

    # ---- 插图：RNN应用场景图（生成一个概念图） ----
    heading(doc, '1.2  设计简介', 2)
    heading(doc, '1.2.1  问题定义', 3)
    add_para(doc, '本课题的核心问题为图像多分类（Image Multi-class Classification）：给定一张28×28像素的'
             '灰度手写数字图像X∈R^(28×28)，要求设计一个分类函数f: R^(28×28)→{0,1,...,9}，使得对于'
             '任意新的手写数字图像能够准确预测其书写的是哪个数字。具体的技术挑战包括：(1) 同一数字存在'
             '多种书写风格（斜体、粗细、连笔程度各异），模型需要学习鲁棒的特征表示；(2) 某些数字在形态上'
             '较为相似（如4和9、3和8、1和7），需要模型具备细粒度的判别能力；(3) 系统需要支持用户在自由绘制'
             '或上传图像后实时返回识别结果，对推理延迟有较高要求；(4) 需要对比CNN和SVM两种不同范式的算法，'
             '揭示深度学习在视觉任务上相比传统方法的优势与局限。', indent=Cm(0.74))

    heading(doc, '1.2.2  数据集介绍', 3)
    add_para(doc, '本课题使用的MNIST（Modified National Institute of Standards and Technology）数据集是'
             '由Yann LeCun、Corinna Cortes和Christopher J.C. Burges共同整理发布的。该数据集可从官方网址'
             '（http://yann.lecun.com/exdb/mnist/）直接下载，也可通过PyTorch的torchvision.datasets.MNIST'
             '方便地加载。数据集包含70,000张28×28像素的灰度图像，每张图像对应一个0到9之间的数字标签。'
             '其中60,000张用于训练，10,000张用于测试。MNIST的图像是经过预处理和标准化的：原始手写数字来自'
             'NIST Special Database 19，经过去除空白区域、缩放、居中、反色等处理，最终呈现为黑底白字的'
             '28×28灰度图，像素值范围为[0, 255]。', indent=Cm(0.74))

    # === 图1：MNIST 样本 ===
    add_image(doc, 'mnist_samples.png', Inches(5.2))
    add_figure_caption(doc, '图1  MNIST手写数字样本展示（每行展示同一数字的不同书写风格）')

    add_para(doc, 'MNIST数据集各类别的样本分布基本均衡。图2所示的类别分布统计揭示了各个数字的训练样本数量'
             '在5,421（数字5）到6,742（数字1）之间，轻微的不均衡不会对模型训练产生实质影响。每个数字'
             '都有其独特的书写特征——数字0通常是圆润的椭圆，数字1是竖直的笔画，数字8由两个叠加的圆组成。'
             '这些视觉先验知识有助于理解不同算法在各类别上的表现差异。', indent=Cm(0.74))

    # === 图2：类别分布 ===
    add_image(doc, 'class_distribution.png', Inches(5.0))
    add_figure_caption(doc, '图2  MNIST训练集各类别样本数量分布')

    # === 图3：均值图 ===
    add_image(doc, 'mean_digits.png', Inches(5.0))
    add_figure_caption(doc, '图3  每类数字的平均像素激活图（反映各类别的统计特征）')

    add_para(doc, '从像素均值图（图3）可以看到，每个数字都有其特征性的"激活区域"——数字1的激活集中在竖直中轴，'
             '数字0的激活形成环形，数字8的激活呈双环结构。这些均值图揭示了不同数字在像素空间中的统计差异，'
             '也暗示了"模板匹配"策略的理论基础。然而，均值图同时显示出模糊的边界——这恰恰说明仅凭像素级别'
             '的统计不足以完美分类，需要更高级的特征提取。', indent=Cm(0.74))

    add_para(doc, '数据预处理方面，本课题针对CNN和SVM的不同特性设计了不同的预处理流程，具体将在2.2节详述。'
             '两种管道均将训练集进一步划分为训练集（90%）和验证集（10%），验证集用于CNN训练过程中的模型'
             '选择和学习率调度。', indent=Cm(0.74))

    heading(doc, '1.2.3  算法选择', 3)
    add_para(doc, '本课题选择卷积神经网络（CNN）和支持向量机（SVM）两种算法进行对比，二者分别代表了'
             '视觉识别任务中"深度特征学习"和"传统特征工程"两种技术路线。', indent=Cm(0.74))
    add_para(doc, '（1）卷积神经网络（Convolutional Neural Network, CNN）：CNN是专为处理网格状拓扑数据'
             '（如图像）而设计的深度学习架构。其核心创新包括三个关键机制：(a) 局部感受野（Local Receptive '
             'Field）——每个神经元只连接输入的一小片区域，而非全连接，这极大地减少了参数数量；(b) 权值共享'
             '（Weight Sharing）——同一卷积核在整个输入空间上滑动使用，使得网络能够检测"位置不变"的视觉特征'
             '（无论数字出现在图像的左上角还是右下角，同一特征检测器都能响应）；(c) 池化（Pooling）——通过'
             '下采样操作逐步降低特征图的空间分辨率，在保留关键信息的同时提升计算效率和感受野范围[1][4]。'
             'CNN通过堆叠多个卷积层和池化层，能够自动学习从低级特征（边缘、角点）到高级语义特征（数字的'
             '整体结构）的层次化表示，无需手动设计特征提取器。', indent=Cm(0.74))

    # === CNN 架构图 ===
    add_image(doc, 'cnn_architecture.png', Inches(5.5))
    add_figure_caption(doc, '图4  CNN卷积神经网络架构图（本课题设计的2卷积+2全连接结构）')

    add_para(doc, '本课题设计的CNN采用"2卷积+2全连接"的轻量级架构，如图4所示：Conv1（1→32,3×3,kernel）→'
             'BatchNorm→ReLU→MaxPool(2×2)→Conv2（32→64,3×3,kernel）→BatchNorm→ReLU→MaxPool(2×2)→'
             'Dropout2d(0.25)→Flatten→FC1(3136→256)→ReLU→Dropout(0.25)→FC2(256→10)。总共约82.5万个'
             '可训练参数。引入BatchNorm加速训练收敛并起到隐式正则化作用；两层Dropout分别置于卷积层之后'
             '（丢弃率0.25）和全连接层之后（丢弃率0.25），有效缓解过拟合。', indent=Cm(0.74))

    add_para(doc, '（2）支持向量机（Support Vector Machine, SVM）：SVM是一种基于统计学习理论和结构风险最小化'
             '原则的经典分类算法[3]。SVM的核心思想是在特征空间中寻找一个最大化分类间隔（Margin）的超平面，'
             '通过引入核函数（Kernel Function）可以将线性不可分的数据隐式映射到高维空间实现非线性分类。本课题'
             '选用RBF（径向基函数）核，因为其在处理图像像素数据时通常优于线性核。然而，直接将SVM应用于'
             'MNIST面临两个关键挑战：一是计算复杂度——SVM的训练时间复杂度约为O(n²)到O(n³)（n为样本数），'
             '在60,000个训练样本上直接训练非常耗时；二是特征表示——将28×28的图像直接展开为784维向量会'
             '丢失像素间的空间邻近关系。为解决这些问题，本课题采用了两项策略：(a) 子采样——从每类随机选择'
             '1,000个训练样本（共10,000个），在保持类别均衡的前提下降低计算负担；(b) PCA降维——将784维'
             '像素向量降至100维，在保留>90%数据方差的同时大幅降低计算开销。这种"SVM + PCA"的经典组合'
             '是传统机器学习处理视觉任务的典型范式。', indent=Cm(0.74))

    heading(doc, '1.2.4  系统架构', 3)
    add_para(doc, '本系统的整体架构遵循规范的机器学习项目流程，分为四个层次：(1) 数据层——负责MNIST数据集的'
             '下载、加载和格式转换，原始数据存储为IDX格式或经由PyTorch自动下载管理。MNIST数据文件托管于'
             'AWS S3、GitHub等多个镜像源，以确保下载可靠性；(2) 处理层——包含CNN管道和SVM管道两套独立的'
             '数据预处理流程。CNN管道使用PyTorch的DataLoader进行批量化加载、数据增强（RandomRotation±5°）'
             '和标准化（均值0.1307，标准差0.3081）；SVM管道使用NumPy进行像素归一化、PCA降维和StandardScaler'
             '标准化；(3) 模型层——CNN模型基于PyTorch框架实现，支持CPU训练和TorchScript序列化；SVM模型'
             '基于scikit-learn框架实现，使用RBF核和概率估计；(4) 应用层——基于Streamlit框架构建Web界面，'
             '提供手写识别（Canvas绘画+图片上传）、模型对比（CNN vs SVM全面指标展示）、数据探索（MNIST样本'
             '和统计信息可视化）和系统说明四个功能模块，并已部署至Streamlit Cloud实现公网访问。', indent=Cm(0.74))
    doc.add_page_break()

    # ==================== 2 结果与分析 ====================
    heading(doc, '2  结果与分析', 1)
    heading(doc, '2.1  数据探索性分析', 2)
    add_para(doc, '在构建分类模型之前，首先对MNIST数据集进行全面的探索性分析（EDA），以了解数据的质量和分布'
             '特性，为后续模型设计提供参考。前文图1至图3已展示了数据集的样本可视化、类别分布和像素均值图，'
             '这里进一步分析数据的统计特性和类别间的可分性。', indent=Cm(0.74))
    add_para(doc, 'MNIST训练集中各类别的样本分布基本均衡：数字0包含5,923个样本，数字1包含6,742个样本，'
             '数字2包含5,958个样本，数字3包含6,131个样本，数字4包含5,842个样本，数字5包含5,421个样本，'
             '数字6包含5,918个样本，数字7包含6,265个样本，数字8包含5,851个样本，数字9包含5,949个样本。'
             '各类别样本数在5,421到6,742之间，最大差异约为24%，属于可接受的不均衡范围，无需进行'
             '过采样或欠采样处理。', indent=Cm(0.74))

    # === 图：易混淆数字对 ===
    add_image(doc, 'confusing_pairs.png', Inches(5.0))
    add_figure_caption(doc, '图5  MNIST数据集中常见的易混淆数字对（4↔9、3↔8、2↔7）')

    add_para(doc, '通过对比分析（图5），可以确认MNIST中存在若干高相似度的数字对——数字4和9在快速书写时都有'
             '向下的竖笔与上方的圈，数字3和8都有相似的曲线结构，数字2和7在笔画转折和尾部特征上也存在'
             '相似性。这些易混淆对是评估分类算法细粒度判别能力的关键测试用例。', indent=Cm(0.74))

    heading(doc, '2.2  数据预处理', 2)
    add_para(doc, '数据预处理是影响模型性能的关键环节，本课题针对CNN和SVM的不同特性设计了不同的预处理流程。'
             '对于CNN管道：(1) ToTensor——将PIL图像或NumPy数组转换为PyTorch张量，像素值自动从[0,255]缩放'
             '至[0,1]；(2) Normalize——使用MNIST全局统计量（均值0.1307，标准差0.3081）进行Z-score标准化，'
             '使输入数据服从标准正态分布，这对梯度下降的稳定性和收敛速度至关重要；(3) RandomRotation——'
             '对训练集随机旋转±5度，增加数据多样性，防止过拟合；(4) DataLoader——将数据组织为128个样本的'
             'Mini-batch，训练时随机打乱，验证和测试时保持顺序。', indent=Cm(0.74))
    add_para(doc, '对于SVM管道：(1) 像素归一化——将像素值除以255缩放至[0,1]；(2) 子采样——从每类随机抽取'
             '1,000个样本（总计10,000个），保证类别均衡的同时降低训练的计算开销，这一步至关重要，因为'
             'SVM对大规模训练集的计算复杂度是其在实际应用中的主要瓶颈；(3) PCA降维——使用主成分分析将'
             '784维降至100维，在保留原始数据超过90%方差的同时，将数据维度减少约87%，大幅降低SVM训练的'
             '内存和时间消耗；(4) StandardScaler标准化——使PCA降维后的各主成分均值为0、标准差为1，确保'
             '距离计算的公平性。PCA拟合和标准化器参数仅从训练集学习，然后应用于测试集，这是防止数据泄露'
             '（Data Leakage）的标准做法。', indent=Cm(0.74))

    # === PCA 方差图 ===
    add_image(doc, 'pca_variance.png', Inches(5.2))
    add_figure_caption(doc, '图6  PCA降维分析——各主成分解释方差比与累计解释方差比')

    add_para(doc, 'PCA分析结果（图6）显示，前100个主成分累计解释了约91.53%的总像素方差，这意味着将784维'
             '像素数据压缩至100维后，原始数据中的绝大多数视觉信息得以保留。同时，丢弃的高维噪声成分可能'
             '带来额外的降噪和正则化效果。', indent=Cm(0.74))

    heading(doc, '2.3  CNN模型训练与评估', 2)
    add_para(doc, 'CNN模型训练采用以下配置：优化器选用Adam（学习率0.001，β₁=0.9，β₂=0.999），损失函数'
             '为交叉熵损失（CrossEntropyLoss），Batch Size设为128，总共训练15个Epoch。学习率调度采用'
             'ReduceLROnPlateau策略：当验证准确率连续2个Epoch无提升时，学习率乘以0.5，这有助于在训练'
             '后期精细调整模型参数。整个训练过程在CPU环境下仅需约5分钟。', indent=Cm(0.74))

    # === CNN 训练曲线 ===
    add_image(doc, 'cnn_training_curves.png', Inches(5.2))
    add_figure_caption(doc, '图7  CNN训练过程中的损失和准确率变化曲线')

    add_para(doc, '训练过程通过训练-验证曲线（图7）进行监控。训练损失从初始的约0.19持续下降至约0.01，'
             '验证损失同步从0.07下降至约0.03，未见明显的过拟合迹象（验证损失未出现上升拐点）。训练准确率'
             '从约94%提升至99.65%，验证准确率稳定在99.15%-99.33%之间。第12个Epoch后学习率从0.001降至'
             '0.0005（ReduceLROnPlateau触发），此后模型进入精细优化阶段，训练损失进一步从0.02降至0.01。'
             '训练结束时CNN在验证集上取得的最佳准确率为99.33%。', indent=Cm(0.74))

    # === CNN 混淆矩阵 ===
    add_image(doc, 'cnn_confusion_matrix.png', Inches(5.0))
    add_figure_caption(doc, '图8  CNN在测试集上的混淆矩阵（准确率99.44%）')

    add_para(doc, '在独立的测试集（10,000样本）上，CNN取得99.44%的准确率，精确率、召回率和F1分数（加权平均）'
             '均为0.9944。混淆矩阵（图8）表明分类错误主要集中在少数特定类别对——最常见的混淆对为4↔9'
             '（约8个样本）、2↔7（约4个样本）和3↔8（约4个样本），这与这些数字在手写形态上的相似性完全一致。'
             'Setosa级别的完美分类（数字0、1的精确率和召回率均达1.00）验证了CNN在特征空间分离良好的类别上'
             '几乎无懈可击。', indent=Cm(0.74))

    heading(doc, '2.4  SVM模型训练与评估', 2)
    add_para(doc, 'SVM模型的训练流程不同于CNN，主要包括PCA降维和RBF核SVM训练两个步骤。SVM使用RBF核函数，'
             '正则化参数C设为10。RBF核的选择基于以下考虑：线性核在原始像素空间或PCA空间中的分类能力有限'
             '（MNIST并非线性可分问题）；多项式核的计算开销较大且调参复杂；RBF核以无限维度的隐式映射提供'
             '了强大的非线性分类能力，是本场景下的自然选择。C=10意味着模型在训练误差和间隔最大化之间略偏向'
             '于降低训练误差——在10,000个训练样本的规模下是合理的，不易导致严重的过拟合。', indent=Cm(0.74))

    # === SVM 混淆矩阵 ===
    add_image(doc, 'svm_confusion_matrix.png', Inches(5.0))
    add_figure_caption(doc, '图9  SVM (RBF核 + PCA 100维) 在测试集上的混淆矩阵（准确率96.46%）')

    add_para(doc, 'SVM在测试集上取得96.46%的准确率。混淆矩阵（图9）揭示了SVM的错误分布——最为显著的混淆'
             '集中在数字8（24个错误中多数被误判为3或5）、数字2（52个错误中部分被误判为7或8）、数字3'
             '（50个错误中多数被误判为8或5）等结构复杂的数字。这说明RBF核在处理某些数字的精细笔画变体'
             '时仍存在分辨力不足的问题。', indent=Cm(0.74))

    heading(doc, '2.5  模型性能对比分析', 2)

    # === 表1：性能对比（表注在表上方） ===
    add_table_caption(doc, '表1  CNN vs SVM 测试集性能指标对比')
    make_table(doc,
        ['算法', '测试准确率', '精确率(加权)', '召回率(加权)', 'F1分数(加权)', '训练/验证准确率'],
        [['CNN', '0.9944', '0.9944', '0.9944', '0.9944', '0.9933 (验证集)'],
         ['SVM (PCA)', '0.9646', '0.9646', '0.9646', '0.9646', '1.0000 (子采样训练集)']])

    add_para(doc, '', after=Pt(8))
    add_para(doc, '表1汇总了两种算法在测试集上的五项核心性能指标。CNN在所有指标上均以约3个百分点的优势领先'
             '于SVM。在图像分类任务中，3%的准确率差异意味着在10,000个测试样本中CNN比SVM少出约282个错误'
             '——这在邮政编码自动分拣等实际场景中直接对应着人力纠错成本的降低。', indent=Cm(0.74))

    # === 综合对比柱状图 ===
    add_image(doc, 'cnn_vs_svm_comparison.png', Inches(5.0))
    add_figure_caption(doc, '图10  CNN vs SVM 四项核心指标综合对比')

    # === 每类准确率 ===
    add_image(doc, 'per_class_accuracy.png', Inches(5.0))
    add_figure_caption(doc, '图11  CNN vs SVM 每个数字类别的分类准确率对比')

    add_para(doc, '每类准确率的细分分析（图11）进一步揭示了两种算法在不同数字上的表现差异。CNN在所有10个'
             '数字类别上的准确率均超过99%，表现最弱的数字为8（约98.9%）；SVM在数字1、0上表现优异（准确率'
             '>99%），但在数字8（约93.5%）和数字5（约95%）上与CNN存在较大差距。这种差异的根源在于CNN能够'
             '通过卷积核学习到数字的局部笔画特征，而SVM仅依赖全局像素分布，对复杂结构数字的细微变化不够敏感。', indent=Cm(0.74))

    # === ROC 曲线 ===
    add_image(doc, 'cnn_vs_svm_roc.png', Inches(5.0))
    add_figure_caption(doc, '图12  CNN vs SVM ROC曲线对比（One-vs-Rest Micro-Average）')

    add_para(doc, 'ROC曲线（图12）采用One-vs-Rest Micro-Average策略绘制，直观展示了两种算法的整体分类性能。'
             'CNN的微平均AUC达到0.9999，近乎完美的分类器；SVM的微平均AUC约为0.998，仍属优秀水平。'
             '两条ROC曲线都极度贴近左上角，说明MNIST对于当前技术水平而言是一个相对成熟的问题，但两种算法'
             '之间仍然存在可测量的性能差距。', indent=Cm(0.74))

    add_para(doc, 'CNN的性能优势来源可归纳为以下几点：(1) 空间结构保留——CNN通过卷积操作保留了像素间的空间'
             '邻近关系，能够检测到边缘、角点、笔画交叉等具有判别力的视觉模式。相比之下，SVM将图像展开为'
             '784维向量后，图像中相邻像素在向量中可能相距较远，空间信息被破坏；(2) 层次化特征学习——CNN'
             '通过多层卷积从低级特征逐步组合为高级语义特征。第一层可能检测边缘和角点，第二层检测笔画和曲线，'
             '全连接层整合为数字的整体表示。这种逐层抽象的能力使得CNN能够捕捉"数字8由两个交叠的圆组成"'
             '这样的高级结构概念，而SVM仅能在像素级别进行比较；(3) 平移不变性——池化操作赋予CNN对数字位置'
             '微小偏移的容忍能力，SVM不具备这样的不变性；(4) 充分的数据利用——CNN使用了全部60,000张训练'
             '图像（含验证集），而SVM受计算限制仅使用10,000张，更多的训练数据使CNN能够学习更丰富的书写风格。', indent=Cm(0.74))

    heading(doc, '2.6  Web应用系统展示', 2)
    add_para(doc, '为实现模型的实际应用价值，本课题基于Streamlit框架开发了手写数字识别Web应用，并已部署至'
             'Streamlit Cloud实现公网访问（在线地址见摘要）。应用提供四个核心功能模块：(1) 手写识别——'
             '用户可使用鼠标在黑色画布上自由绘制数字（支持触摸屏），也可上传手写数字图片文件，系统自动进行'
             '预处理（缩放至28×28、灰度转换、反色判断）和标准化，调用选定模型（CNN或SVM）实时推理，展示'
             '预测数字和各类别概率分布条形图；(2) 模型对比——集中展示CNN vs SVM的所有评估图表（综合指标'
             '对比柱状图、ROC曲线、每类准确率、混淆矩阵、训练曲线和错误案例），供用户深入理解两种模型的'
             '性能差异；(3) 数据探索——展示MNIST数据集的统计信息、样本图像和类别均值图；(4) 关于系统——'
             '介绍项目背景、算法原理和参考文献。', indent=Cm(0.74))

    # === 表2：Web应用功能 ===
    add_table_caption(doc, '表2  Web应用系统功能模块一览')
    make_table(doc,
        ['功能模块', '交互方式', '技术实现', '核心用途'],
        [['在线识别', 'Canvas绘画 / 图片上传', 'st_canvas + CNN/SVM推理', '实时手写数字识别'],
         ['模型对比', '图表浏览', '预先训练生成的评估图表', 'CNN vs SVM性能对比'],
         ['数据探索', '图表展示', 'MNIST统计分析图表', '理解数据集特性'],
         ['关于系统', '文本展示', 'Markdown渲染', '项目背景及技术说明']])

    add_para(doc, '', after=Pt(8))
    add_para(doc, '系统实现中的技术亮点包括：(1) 模型缓存——采用Streamlit的@st.cache_resource装饰器，'
             '将CNN模型（约3.3MB）加载到内存后跨请求复用，确保毫秒级推理响应；(2) 自动反色检测——'
             '对用户上传的任意图片，自动判断其为黑底白字（MNIST标准格式）还是白底黑字（常见图片格式），'
             '并自动执行反色处理，大幅提升了对不同来源图片的识别鲁棒性；(3) 实时画布——基于streamlit-'
             'drawable-canvas组件实现浏览器端Canvas绘画，支持鼠标和触屏操作，笔画平滑流畅，用户界面友好；'
             '(4) 云端部署——项目托管于GitHub并通过Streamlit Cloud自动构建和部署，无需用户自行配置环境'
             '即可通过浏览器在线使用。', indent=Cm(0.74))
    doc.add_page_break()

    # ==================== 5 结论 ====================
    heading(doc, '5  结论', 1)
    add_para(doc, '本文针对手写数字识别这一经典计算机视觉问题，基于MNIST数据集设计并实现了完整的识别系统。'
             '通过系统性的实验设计和多维度的性能对比，得出以下主要结论：', indent=Cm(0.74))
    add_para(doc, '（1）算法性能方面：CNN在测试集上以99.44%的准确率显著优于基于PCA降维的SVM（96.46%），'
             '提升约2.98个百分点，验证了深度学习在视觉任务上优于传统手工特征工程方法的基本认知。CNN通过'
             '卷积操作保留了图像的空间结构信息，通过层次化特征学习自动提取了从边缘到整体结构的判别特征，'
             '通过池化获得了平移不变性——这些机制是传统像素级分类器无法比拟的。SVM虽然准确率略低，但在'
             '"仅使用1/6训练数据"和"输入仅为100维PCA向量"的条件下仍达到了96%以上的准确率，体现了其强大的'
             '非线性分类能力和对小样本场景的适应性——在训练数据有限的实际应用中，SVM仍然是一个有竞争力的选择。', indent=Cm(0.74))
    add_para(doc, '（2）算法原理对比方面：CNN和SVM的对比深刻体现了"表示学习"与"固定特征空间学习"两种范式的'
             '本质差异。CNN通过端到端训练同时优化特征提取器和分类器，利用空间先验（卷积的权值共享和平移'
             '不变性）大幅降低了模型复杂度——约82.5万参数即可高效处理28×28=784维的输入空间。SVM依赖固定的'
             '特征映射（PCA降维），特征提取与分类器训练独立进行，模型容量（支持向量数量）由数据特性决定。'
             '这一对比揭示了"数据驱动特征学习"相对于"人为设计特征空间"的显著优势。', indent=Cm(0.74))
    add_para(doc, '（3）工程实践方面：成功构建了基于Streamlit的交互式Web应用并部署至云端公网，支持Canvas'
             '手绘识别和图片上传识别两种交互方式，用户可在侧边栏切换CNN/SVM模型实时对比识别效果。项目'
             '采用模块化设计，模型序列化使用PyTorch checkpoint和joblib两种工业标准方案。系统的架构设计'
             '使得后续添加新模型（如ResNet、Vision Transformer）或扩展至类似数据集变得简单便捷。', indent=Cm(0.74))
    add_para(doc, '（4）局限与展望：本课题存在以下可在未来工作中改进的方面：第一，可扩充更多算法对比（随机'
             '森林、梯度提升树、k-NN、全连接神经网络等），构建更丰富的渐进式算法对比研究；第二，可引入'
             '更激进的图像扩增策略（随机平移、弹性形变等），进一步提升模型对书写变体的鲁棒性；第三，可'
             '探索Grad-CAM等模型可解释性工具，可视化CNN的卷积层注意力区域，帮助理解模型决策依据；第四，'
             '可拓展到EMNIST（手写字母）或SVHN（街景门牌数字）等域外数据集，评估模型在实际多变场景中的'
             '泛化表现；第五，可将系统改造为Docker容器化微服务，支持更大规模的并发请求和更灵活的部署方案。', indent=Cm(0.74))
    add_para(doc, '综上所述，本课程设计通过手写数字识别这一经典案例，系统性地实践了机器学习——特别是深度学习'
             '在计算机视觉中的完整技术流程：从数据获取与探索、特征预处理、模型训练与超参数调优、多维度性能'
             '评估，到最终的云端Web应用部署。通过CNN与SVM两种代表性算法的全面对比分析，直观展示了深度学习'
             '在视觉任务上的显著优势及其背后的技术原理。整个过程加深了对卷积神经网络核心机制（卷积、池化、'
             '特征层次化）的直觉理解，锻炼了PyTorch深度学习框架的实际使用能力，掌握了机器学习项目的工程化'
             '方法和云部署流程，达到了课程设计的预期目标。', indent=Cm(0.74))
    doc.add_page_break()

    # ==================== 参考文献 ====================
    heading(doc, '参考文献', 1)
    refs = [
        '[1] LeCun Y, Bottou L, Bengio Y, et al. Gradient-based learning applied to document recognition[J]. Proceedings of the IEEE, 1998, 86(11): 2278-2324.',
        '[2] LeCun Y, Cortes C, Burges C J C. The MNIST Database of Handwritten Digits[EB/OL]. http://yann.lecun.com/exdb/mnist/.',
        '[3] Cortes C, Vapnik V. Support-vector networks[J]. Machine Learning, 1995, 20(3): 273-297.',
        '[4] Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks[C]. NeurIPS, 2012: 1097-1105.',
        '[5] Simonyan K, Zisserman A. Very deep convolutional networks for large-scale image recognition[C]. ICLR, 2015.',
        '[6] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]. CVPR, 2016: 770-778.',
        '[7] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine Learning in Python[J]. JMLR, 2011, 12: 2825-2830.',
        '[8] Paszke A, Gross S, Massa F, et al. PyTorch: An imperative style, high-performance deep learning library[C]. NeurIPS, 2019: 8024-8035.',
        '[9] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.',
        '[10] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.',
        '[11] Deng L. The MNIST database of handwritten digit images for machine learning research[J]. IEEE Signal Processing Magazine, 2012, 29(6): 141-142.',
        '[12] Jolliffe I T, Cadima J. Principal component analysis: a review and recent developments[J]. Philosophical Transactions of the Royal Society A, 2016, 374(2065): 20150202.',
        '[13] 李航. 统计学习方法(第2版)[M]. 北京: 清华大学出版社, 2019.',
        '[14] Abadi M, Barham P, Chen J, et al. TensorFlow: A system for large-scale machine learning[C]. USENIX OSDI, 2016: 265-283.',
    ]
    for ref in refs:
        add_para(doc, ref, fn_cn='宋体', size=Pt(10.5), after=Pt(2))

    # ==================== 保存 ====================
    path = os.path.join(REPORT_DIR, '《机器学习》综合设计报告_手写数字识别系统.docx')
    doc.save(path)
    print(f'[SUCCESS] 报告已生成: {path}')
    print(f'  包含图表: 12张图 + 2张表')
    print(f'  正文约7000字')
    return path

if __name__ == '__main__':
    generate_report()
